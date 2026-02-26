#!/usr/bin/env python3
"""Word (docx) to HTML 5 converter with CLI and API function support."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def convert_docx_to_html(docx_path: str | Path) -> str:
    """Convert a Word document to HTML 5.
    
    Args:
        docx_path: Path to the .docx file.
        
    Returns:
        HTML 5 string representation of the document.
        
    Raises:
        FileNotFoundError: If the .docx file doesn't exist.
        ValueError: If the file is not a valid .docx file.
    """
    path = Path(docx_path)
    
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx file, got: {path.suffix}")
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    
    doc = Document(path)
    html_parts = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="UTF-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        "<title>Converted Document</title>",
        "<style>",
        "body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }",
        "h1, h2, h3, h4, h5, h6 { margin-top: 1em; margin-bottom: 0.5em; }",
        "table { border-collapse: collapse; width: 100%; margin: 1em 0; }",
        "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
        "th { background-color: #f2f2f2; }",
        "ul, ol { margin: 1em 0; padding-left: 40px; }",
        "li { margin: 0.25em 0; }",
        "blockquote { border-left: 4px solid #ccc; margin: 1em 0; padding-left: 1em; color: #666; }",
        "pre { background: #f4f4f4; padding: 1em; overflow-x: auto; }",
        "</style>",
        "</head>",
        "<body>",
    ]
    
    para_idx = 0
    table_idx = 0
    elements = list(doc.element.body)
    
    for element in elements:
        if element.tag.endswith('p'):
            if para_idx < len(doc.paragraphs):
                html_parts.append(_convert_paragraph(doc.paragraphs[para_idx]))
                para_idx += 1
        elif element.tag.endswith('tbl'):
            if table_idx < len(doc.tables):
                html_parts.append(_convert_table(doc.tables[table_idx]))
                table_idx += 1
    
    html_parts.append("</body>")
    html_parts.append("</html>")
    
    return "\n".join(html_parts)


def _convert_paragraph(paragraph: Paragraph) -> str:
    """Convert a single paragraph to HTML."""
    if not paragraph.text.strip():
        return ""
    
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    
    if style_name.startswith('heading'):
        level = style_name.replace('heading ', '')
        try:
            level = int(level)
        except ValueError:
            level = 1
        level = min(max(level, 1), 6)
        return f"<h{level}>{_convert_runs(paragraph)}</h{level}>"
    
    if style_name == 'title':
        return f"<h1>{_convert_runs(paragraph)}</h1>"
    
    if paragraph.style.name == 'List Bullet' or paragraph.style.name == 'List Bullet HTML':
        return f"<ul><li>{_convert_runs(paragraph)}</li></ul>"
    
    if paragraph.style.name == 'List Number' or paragraph.style.name == 'List Number HTML':
        return f"<ol><li>{_convert_runs(paragraph)}</ol>"
    
    if style_name == 'quote' or style_name == 'blockquote':
        return f"<blockquote>{_convert_runs(paragraph)}</blockquote>"
    
    return f"<p>{_convert_runs(paragraph)}</p>"


def _convert_runs(paragraph: Paragraph) -> str:
    """Convert runs within a paragraph, preserving formatting."""
    result = ""
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        
        formatted = _convert_run(run)
        result += formatted
    
    return result


def _convert_run(run: Run) -> str:
    """Convert a single run with its formatting."""
    text = run.text
    
    if not text:
        return ""
    
    tags = []
    closing_tags = []
    
    if run.bold:
        tags.append("strong")
        closing_tags.append("strong")
    
    if run.italic:
        tags.append("em")
        closing_tags.append("em")
    
    if run.underline:
        tags.append("u")
        closing_tags.append("u")
    
    if run.style and 'strike' in run.style.name.lower():
        tags.append("s")
        closing_tags.append("s")
    
    if not tags:
        return text
    
    html = text
    for tag in reversed(tags):
        html = f"<{tag}>{html}</{tag}>"
    
    return html


def _convert_table(table: Table) -> str:
    """Convert a table to HTML."""
    if not table.rows:
        return ""
    
    html = ["<table>"]
    
    for row_idx, row in enumerate(table.rows):
        html.append("<tr>")
        for cell in row.cells:
            cell_text = cell.text.strip()
            if row_idx == 0:
                html.append(f"<th>{cell_text}</th>")
            else:
                html.append(f"<td>{cell_text}</td>")
        html.append("</tr>")
    
    html.append("</table>")
    
    return "\n".join(html)


def main(argv: Optional[list[str]] = None) -> int:
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Convert Word (.docx) documents to HTML 5."
    )
    parser.add_argument(
        "input",
        type=Path,
        help="Path to the input .docx file",
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        help="Path to save the output HTML file (default: stdout)",
    )
    parser.add_argument(
        "--pretty",
        action="store_true",
        help="Format output with indentation",
    )
    
    args = parser.parse_args(argv)
    
    try:
        html = convert_docx_to_html(args.input)
        
        if args.pretty:
            html = re.sub(r'>(<)', r'>\n\1', html)
            html = re.sub(r'([a-z]>)', r'\1\n', html)
            html = html.strip()
        
        if args.output:
            args.output.write_text(html)
            print(f"HTML saved to: {args.output}")
        else:
            print(html)
        
        return 0
        
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"Unexpected error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
